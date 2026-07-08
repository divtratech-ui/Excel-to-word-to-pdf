import os
import subprocess
import openpyxl
from docx import Document
from docx2pdf import convert

# ---------- CONFIG ----------
SPREADSHEET = r"Spreadsheet_path"
SHEET_NAME = "Sheet1"               # or None for the active sheet
FIRST_NAME_COL = "First Name"       # column header for first name
LAST_NAME_COL = "Last Name"         # column header for last name
TEMPLATE_DOCX = r"Template_docx_path"  # contains {NAME}
OUTPUT_DIR = "output_pdfs"          # where PDFs will be saved
# ----------------------------

def replace_placeholder(doc, placeholder, replacement):
    """Replace placeholder in paragraphs and tables."""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)


def main():
    # Load workbook and sheet
    wb = openpyxl.load_workbook(SPREADSHEET)
    ws = wb[SHEET_NAME] if SHEET_NAME else wb.active

    # Find column indices by header
    first_col = None
    last_col = None
    for cell in ws[1]:  # first row (headers)
        if cell.value:
            header = cell.value.strip()
            if header == FIRST_NAME_COL:
                first_col = cell.column
            elif header == LAST_NAME_COL:
                last_col = cell.column

    if first_col is None:
        raise ValueError(f"Column '{FIRST_NAME_COL}' not found in row 1.")
    if last_col is None:
        raise ValueError(f"Column '{LAST_NAME_COL}' not found in row 1.")

    # Create output folder
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    # Process each row (skip header)
    for row in ws.iter_rows(min_row=2, values_only=True):
        first_name = row[first_col - 1]   # zero-index adjustment
        last_name = row[last_col - 1]

        # Skip if either name is missing
        if not first_name or not last_name:
            print(f"Skipping row: missing first or last name ({first_name}, {last_name})")
            continue

        first_name = str(first_name).strip()
        last_name = str(last_name).strip()
        full_name = f"{first_name} {last_name}"
        username = f"{first_name}-{last_name}"   # hyphen between names, no spaces
        print(f"Processing: {full_name} (username: {username})")

        # Load template and replace placeholder
        doc = Document(TEMPLATE_DOCX)
        replace_placeholder(doc, "{NAME}", full_name)
        replace_placeholder(doc, "{USERNAME}", username)

        # Save temporary docx (named with full name)
        temp_docx = os.path.join(OUTPUT_DIR, f"{full_name}.docx")
        doc.save(temp_docx)

        # Convert to PDF (outputs {full_name}.pdf in OUTPUT_DIR)
        pdf_path = os.path.join(OUTPUT_DIR, f"{full_name}.pdf")

        # perform conversion for this file
        convert(temp_docx, pdf_path)


if __name__ == "__main__":
    main()